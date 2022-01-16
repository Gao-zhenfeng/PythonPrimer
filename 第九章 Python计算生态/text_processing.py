# -*- coding: utf-8 -*-
"""
-------------------------------------------------
   File Name：     text_processing
   Description :   PyPDF2
   Author :        Amadeus
   date：          2021/12/23
-------------------------------------------------
   Change Activity:
                   2021/12/23:
-------------------------------------------------
PyPDF2: 用来处理PDF文件的工具集，完全用python实现，非常稳定

"""
# %%
from PyPDF2 import PdfFileReader, PdfFileMerger

merger = PdfFileMerger()
input1 = open('E:\\Program\\Python\\Exercise\\第九章 Python计算生态\\电磁波在分层介质中的传播1.pdf', 'rb')
input2 = open('E:\\Program\\Python\\Exercise\\第九章 Python计算生态\\电磁波在分层介质中的传播2.pdf', 'rb')
merger.append(fileobj=input1, pages=(0, 3))
merger.merge(position=2, fileobj=input2, pages=(0, 1))
output = open('output.pdf', 'wb')
merger.write(output)
merger.close()
output.close()

# %%
'''
NLTK:自然语言文本处理的第三方库
Python-docx:创建或更新word文件的第三方库
'''

from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')


